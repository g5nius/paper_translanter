import re
import docx
from io import StringIO

from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser


def get_token():
    url = "https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id=&client_secret="

    payload = ""
    headers = {
        'Content-Type': 'application/json',
        'Accept': 'application/json'
    }

    response = requests.request("POST", url, headers=headers, data=payload)

    print(response.text)



output_string = StringIO()
with open('smote.synthetic minority oversampling technique.pdf', 'rb') as in_file:
    parser = PDFParser(in_file)
    doc = PDFDocument(parser)
    rsrcmgr = PDFResourceManager()
    device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    for page in PDFPage.create_pages(doc):
        interpreter.process_page(page)


output_string.seek(0)

# 处理换行符带来的分割
txt_paragraph = ''
txt_paragraph_list = []
while True:
    s = output_string.readline()
    if s != '\n' and s != '':
        tmp_str = s[:-1]
        txt_paragraph = txt_paragraph + ' ' + tmp_str if txt_paragraph else txt_paragraph + tmp_str
    else:
        txt_paragraph_list.append(txt_paragraph)
        txt_paragraph = ''
    if s == '':
        break


# 处理段落是否被页眉页脚插图中断
# 删减邮箱无用信息
# 删去引文及以后部分 TODO:将引文以原格式添加到结尾
# TODO 添加根据文本位置和内容进行机器学习分类
paragraphs_translated = []
del_list = []
new_list = []
for index, paragraph in enumerate(txt_paragraph_list):
    if 'References' == paragraph:
        break
    # email_reg = re.compile(r'[\w]+(\.[\w]+)*@[\w]+(\.[\w])+')
    # if email_reg.search(paragraph):
    #     txt_paragraph_list.pop(index)
    if len(paragraph.split()) < 24:
        del_list.append(paragraph)
    else:
        new_list.append(paragraph)

# 合并两段
# 翻译
# 保存到word文档
import requests

token = ''
url = 'https://aip.baidubce.com/rpc/2.0/mt/texttrans/v1?access_token=' + token


from_lang = 'en'  # example: en
to_lang = 'zh'  # example: zh
term_ids = ''  # 术语库id，多个逗号隔开
headers = {'Content-Type': 'application/json'}
document = docx.Document()
add_flag = 0
add_p = ''
for p in new_list:
    if p[-1] != '.' and p[-1] != ':':
        add_p = p
        continue
    if add_flag:
        p = add_p + p
        add_flag = 0
    else:
        payload = {'q': p, 'from': from_lang, 'to': to_lang, 'termIds': term_ids}
        r = requests.post(url, params=payload, headers=headers)
        result = r.json()
        # print(json.dumps(result, indent=4, ensure_ascii=False))
        try:
            result = result['result']['trans_result'][0]['dst']
        except:
            try:
                r = requests.post(url, params=payload, headers=headers)
                result = r.json()
                result = result['result']['trans_result'][0]['dst']
            except:
                print('error')
        document.add_paragraph(result)
document.save('translated.docx')
